#!/usr/bin/env python3

import configparser
import csv
import httplib2
import io
import os
import re
import sys
import tempfile
from datetime import datetime

import gspread
import oauth2client
from apiclient import discovery, errors
from apiclient.http import MediaFileUpload, MediaIoBaseDownload
from docx import Document
from docxtpl import DocxTemplate
from gspread.exceptions import SpreadsheetNotFound
from jinja2 import Environment
from oauth2client import client
from oauth2client import tools
from openpyxl import load_workbook, Workbook

try:
    import argparse
    flags = argparse.ArgumentParser(parents=[tools.argparser]).parse_args()
except ImportError:
    flags = None


SCOPES = ['https://www.googleapis.com/auth/drive',
          'https://spreadsheets.google.com/feeds']
CLIENT_SECRET_FILE = 'client_secret.json'
APPLICATION_NAME = 'autotemplation'
TEMPLATE_REGEX = r"\{\{ ([A-Za-z0-9_]+) \}\}"
DOCUMENT_TYPE = 'vnd.openxmlformats-officedocument.wordprocessingml.document'
SHEET_TYPE = 'vnd.openxmlformats-officedocument.spreadsheetml.sheet'


def get_credentials():
    """Gets valid user credentials from storage.

    If nothing has been stored, or if the stored credentials are invalid,
    the OAuth2 flow is completed to obtain the new credentials.

    Returns:
        Credentials, the obtained credential.
    """
    home_dir = os.path.expanduser('~')
    credential_dir = os.path.join(home_dir, '.credentials')
    if not os.path.exists(credential_dir):
        os.makedirs(credential_dir)
    credential_path = os.path.join(credential_dir,
                                   'drive-python-quickstart.json')

    store = oauth2client.file.Storage(credential_path)
    credentials = store.get()

    if not credentials or credentials.invalid:
        flow = client.flow_from_clientsecrets(CLIENT_SECRET_FILE, SCOPES)
        flow.user_agent = APPLICATION_NAME
        if flags:
            credentials = tools.run_flow(flow, store, flags)
        else:  # Needed only for compatibility with Python 2.6
            credentials = tools.run(flow, store)
        print('Storing credentials to {}'.format(credential_path))
    return credentials


def get_or_create_destination_folder_id(drive_service,
                                        destination_folder_name):
    folder_mime_type = 'application/vnd.google-apps.folder'
    response = drive_service.files().list(
        q="mimeType='{}' and name='{}'".format(folder_mime_type,
                                               destination_folder_name),
        spaces='drive',
        fields='nextPageToken, files(id, name)').execute()
    files = response.get('files', [])
    if files:
        folder_id = files[0].get('id')
    else:
        file_metadata = {
            'name': destination_folder_name, 'mimeType': folder_mime_type
        }
        file = drive_service.files().create(body=file_metadata,
                                            fields='id').execute()
        folder_id = file.get('id')
    return folder_id


def get_files_in_folder(service, folder_id):
    """Print files belonging to a folder.

    Args:
      service: Drive API service instance.
      folder_id: ID of the folder to print files from.
    """
    files = dict()
    page_token = None
    while True:
        try:
            param = {}
            if page_token:
                param['pageToken'] = page_token
            children = service.files().list(
                q="'{}' in parents and trashed=false".format(folder_id),
                spaces='drive',
                fields='nextPageToken, files(id, name, mimeType)',
                pageToken=page_token).execute()
            for child in children.get('files', []):
                files[child['name']] = child
            page_token = children.get('nextPageToken')
            if not page_token:
                break
        except errors.HttpError as error:
            print('An error occurred: {}'.format(error))
            break
    return files


def get_template(service, folder_ids):
    templates = dict()
    for folder_id in folder_ids:
        templates.update(get_files_in_folder(service, folder_id))
    if len(templates) == 0:
        print("No Templates Found. Please check Folder ID. Exiting.")
        sys.exit(1)
    elif len(templates) == 1:
        first = next(iter(templates.items()))
        print("One Template Found. Using {}.".format(first[0]))
        return first
    template_list = sorted(templates.keys())
    print('Please Select a Template:')
    while True:
        for i, name in enumerate(template_list):
            print('{}) {}'.format(i+1, name))
        print('q) Quit')
        choice = input(
            'Selection?  '.format(len(template_list)))
        if choice.lower() in ['q', 'quit', 'exit']:
            print('Exiting.')
            sys.exit(0)
        try:
            file_name = template_list[int(choice)-1]
            file = templates[file_name]
            break
        except (ValueError, IndexError):
            print('Error: Invalid Selection.')
    return file


def get_date_and_set_context(context_dict):
    while True:
        use_date = input("Please enter document date or leave empty for "
                         "today's date (Ex: 20150120):  ")
        if not use_date:
            date_object = datetime.now()
            break
        else:
            try:
                date_object = datetime.strptime(use_date, '%Y%m%d')
                break
            except (AttributeError, ValueError):
                print("Unable to parse date entry. Please use example format.")
                pass
    context_dict['DATE_FULL'] = date_object.strftime('%B %d, %Y')
    context_dict['DATE_FULL_NUM'] = date_object.strftime('%Y%m%d')
    context_dict['DATE_FULL_DASH'] = date_object.strftime('%m-%d-%Y')
    context_dict['DATE_FULL_SLASH'] = date_object.strftime('%m/%d/%Y')
    context_dict['DATE_MONTH'] = date_object.strftime('%B')
    context_dict['DATE_MONTH_NUM'] = date_object.strftime('%m')
    context_dict['DATE_DAY_FULL'] = date_object.strftime('%A')
    context_dict['DATE_DAY_SHORT'] = date_object.strftime('%a')
    context_dict['DATE_DAY_NUM'] = date_object.strftime('%d')
    context_dict['DATE_YEAR'] = date_object.strftime('%Y')
    #  Determine day suffix
    day_num = int(context_dict['DATE_DAY_NUM'])
    if 11 <= day_num <= 13:
        day_suffix = 'th'
    elif day_num % 10 == 1:
        day_suffix = 'st'
    elif day_num % 10 == 2:
        day_suffix = 'nd'
    elif day_num % 10 == 3:
        day_suffix = 'rd'
    else:
        day_suffix = 'th'
    context_dict['DATE_DAY_SUFFIX'] = day_suffix


def get_table_data_for_csv(doc):
    if doc.tables and len(doc.tables) == 1:
        table = doc.tables[0]
        data = [[cell.text for cell in row.cells] for row in table.rows]
    else:
        data = []
    return data


def get_target_name(template_name, context):
    return Environment().from_string(template_name).render(context)


def get_worksheet(credentials):
    gc = gspread.authorize(credentials)
    worksheet = None
    while not worksheet:
        spreadsheetId = input("Please enter ID of Google Sheet to use:  ")
        #  attempts to get first sheet of book for now
        try:
            worksheet = gc.open_by_key(spreadsheetId).get_worksheet(0)
        except SpreadsheetNotFound:
            print("Invalid Google Sheet ID!")
    return worksheet


def get_worksheet_headers(worksheet):
    worksheet_headers = dict()
    worksheet_headers['columns'] = {
        item.lower(): index + 1 for index, item
        in enumerate(worksheet.row_values(1)) if item}
    worksheet_headers['rows'] = {
        item.lower(): index + 1 for index, item
        in enumerate(worksheet.col_values(1)) if item}
    return worksheet_headers


def worksheet_lookup(worksheet, worksheet_headers, var):
    column, row = var.lower().replace('_', ' ').split('  ')
    try:
        column_key = worksheet_headers['columns'][column]
        row_key = worksheet_headers['rows'][row]
    except KeyError:
        print("ERROR: Lookup failed for Column: {}, Row: {}".format(column,
                                                                    row))
        return "{} {} {}".format("{{", var, "}}")

    return worksheet.cell(row_key, column_key).value


def get_sheet_data(file_handler):
    wb = load_workbook(file_handler, read_only=True)
    data = []
    sheet = wb.get_active_sheet()
    for row in sheet.iter_rows():
        data_row = []
        for cell in row:
            data_row += [cell.value]
        # ignore empty rows
        if any(data_row):
            data += [data_row]
    return data


def get_mime_type(google_mime_type):
    if 'document' in google_mime_type:
        mime_type = 'application/{}'.format(DOCUMENT_TYPE)
        sheet = False
    elif 'sheet' in google_mime_type:
        mime_type = 'application/{}'.format(SHEET_TYPE)
        sheet = True
    else:
        raise TypeError("Unknown MIME from Google: {}".format(
            google_mime_type))
    return mime_type, sheet


def get_template_variables(doc, template_name):
    template_vars = set()

    #  Template Name
    template_vars |= set(re.findall(TEMPLATE_REGEX, template_name))
    #  Paragraphs
    template_vars |= set([var for var_list in
                          [re.findall(TEMPLATE_REGEX, p.text)
                           for p in doc.paragraphs] for var in var_list])
    #  Tables
    template_vars |= set([var for var_list in
                          [re.findall(TEMPLATE_REGEX, cell.text)
                           for table in doc.tables
                           for row in table.rows
                           for cell in row.cells] for var in var_list])
    return template_vars


def main():
    config = configparser.ConfigParser()
    config.read('autotemplation.ini')
    template_folder_ids = config['DEFAULT']['TemplateFolderID'].split(',')
    destination_folder_name = config['DEFAULT']['DestinationFolderName']
    credentials = get_credentials()
    http = credentials.authorize(httplib2.Http())
    drive_service = discovery.build('drive', 'v3', http=http)
    destination_folder_id = get_or_create_destination_folder_id(
        drive_service, destination_folder_name)
    template_file = get_template(drive_service, template_folder_ids)
    mime_type, is_sheet = get_mime_type(template_file['mimeType'])
    request = drive_service.files().export_media(
        fileId=template_file['id'],
        mimeType=mime_type)
    fh = io.BytesIO()
    downloader = MediaIoBaseDownload(fh, request)
    done = False
    while done is False:
        status, done = downloader.next_chunk()
        print("Download %d%%." % int(status.progress() * 100))
    if is_sheet:
        print("Spreadsheet selected, converting to Doc. (Slow)")
        table_data = get_sheet_data(fh)
        row_count = len(table_data)
        col_count = len(table_data[0])
        document = Document()
        doc_table = document.add_table(rows=row_count,
                                       cols=col_count)
        for r, row in enumerate(table_data):
            row_cells = doc_table.rows[r].cells
            print("Converting row {}/{}...".format(r+1, row_count), end="\r")
            for i, cell in enumerate(row):
                if cell:
                    row_cells[i].text = cell
        print("Conversion complete. "
              "Warning: Processing large sheets will take some time.")
        temp_doc_file = io.BytesIO()
        document.save(temp_doc_file)
        doc = DocxTemplate(temp_doc_file)
    else:
        doc = DocxTemplate(fh)
    full_doc = doc.get_docx()

    template_vars = get_template_variables(full_doc, template_file['name'])
    if any('__' in x for x in template_vars):
        worksheet = get_worksheet(credentials)
        worksheet_headers = get_worksheet_headers(worksheet)
    context = dict()
    get_date_and_set_context(context)
    for var in template_vars:
        if var not in context:
            if '__' in var:
                context[var] = worksheet_lookup(
                    worksheet, worksheet_headers, var)
            else:
                context[var] = input("Enter a value for {}:  ".format(var))
    new_file_name = get_target_name(template_file['name'], context)
    doc.render(context)
    temp_file = tempfile.NamedTemporaryFile()
    doc.save(temp_file)
    if is_sheet:
        csv_name = '{}.csv'.format(new_file_name)
        doc_csv = DocxTemplate(temp_file)
        csv_data = get_table_data_for_csv(doc_csv)
        if csv_data:
            with open(csv_name, 'w') as output:
                writer = csv.writer(output, lineterminator='\n')
                writer.writerows(csv_data)
            print('{} created in local folder'.format(csv_name))
        else:
            print('Unable to create CSV. '
                  'Less than or more than 1 table found.')
        workbook = Workbook()
        sheet = workbook.get_active_sheet()
        for row in csv_data:
            sheet.append(row)
        workbook.save(temp_file)
        upload_mimetype = 'application/vnd.google-apps.spreadsheet'
    else:
        upload_mimetype = 'application/vnd.google-apps.document'

    file_metadata = {
        'name': new_file_name,
        'parents': [destination_folder_id],
        'mimeType': upload_mimetype
    }
    media = MediaFileUpload(temp_file.name,
                            mimetype=mime_type,
                            resumable=True)
    drive_service.files().create(body=file_metadata,
                                 media_body=media,
                                 fields='id').execute()
    print('{} placed in folder {}.'.format(new_file_name,
                                           destination_folder_name))
    temp_file.close()

if __name__ == '__main__':
    main()
